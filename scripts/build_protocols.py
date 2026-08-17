#!/usr/bin/env python3
# /// script
# requires-python = ">=3.11"
# dependencies = ["python-docx>=1.1"]
# ///
"""Render a project's protocol .docx set and cross-check it against the sample tree.

Usage (from a curation project root):

    uv run --script <PLUGIN>/scripts/build_protocols.py
    uv run --script <PLUGIN>/scripts/build_protocols.py --only AFM --force
    uv run --script <PLUGIN>/scripts/build_protocols.py --coverage-only

Phase 3b's judgment lives in two JSON files the model writes; this script is the
deterministic half. It renders, it verifies, and it reports coverage. It never
decides which manuscript section belongs in which protocol.

Input: protocols/_methods.json (verbatim manuscript excerpts):

    [
      {
        "heading": "Analysis of confocal microscopy images",  required
        "paras": ["...", "..."],                              required, verbatim
        "verbatim": true                                      optional, default true
      }
    ]

A heading may occur more than once (Oak et al. has two "Nanoneedle AFM"
sections). Occurrences are consumed in document order. Set "verbatim": false
only for content that is genuinely NOT a byte-for-byte copy of the source. A
transcribed display equation the PDF floats as un-extractable math is the
motivating case. Everything else must survive the verbatim check below.

Input: protocols/_manifest.json (the mapping, and the only place judgment sits):

    {
      "lab": "SHE",                     required, uppercase lab tag
      "stamp": "260807",                required, YYMMDD batch stamp
      "version": 1,                     optional, default 1  -> V1
      "study": "Oak et al., ...",       optional, printed in COVERAGE.md
      "protocols": [
        {
          "topic": "ConfocalImageAnalysis",         required, CamelCase
          "headings": ["Analysis of confocal ..."], required, order preserved
          "assays": ["Imaging Analysis"],           optional, exact assay titles
          "note": "..."                             optional, carried into Table A
        }
      ]
    }

Filenames are derived, never declared: ``P.<LAB>-<STAMP>-V<n>_<Topic>.docx``.
All 493 stamped SOPs in the MetNet registry use this form.

Input: sample_tree.json (optional): supplies the coverage tables. Edges carry
``from`` / ``to`` / ``assays``; nodes carry ``id`` / ``count``. Absent, the
script still renders the documents and writes a COVERAGE.md that says why the
tables are missing.

Input: protocols/_sops.json (optional): written by upload_sops.py, maps a
filename to its registered SOP id. Fills the SOP column of Table A.

Outputs, all under protocols/:

    P.<LAB>-<STAMP>-V<n>_<Topic>.docx   one per manifest entry
    COVERAGE.md                         Table A + Table B, regenerated every run

An existing .docx is NEVER overwritten without --force. Protocol documents get
registered on a live server and emailed to PIs; silently rewriting one that has
already been handed over is the failure this guard exists to prevent.

House style, copied from the Shenoy and Cui curations: Normal style throughout,
Times New Roman 12pt, US Letter, 1" margins, the manuscript's own section
heading in italics, body text verbatim, and nothing else. No metadata block, no
curation notes, no commentary. That analysis belongs in COVERAGE.md and the
project's protocols/README.md, not inside the document a PI reads.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

FILENAME_RE = re.compile(r"^P\.[A-Z0-9]+-\d{6}-V\d+_[A-Za-z0-9]+\.docx$")
TOPIC_RE = re.compile(r"^[A-Za-z0-9]+$")


class ProtocolError(RuntimeError):
    """A manifest, methods file or tree the curator has to fix."""


# ── inputs ────────────────────────────────────────────────────────────────────

def load_json(path: Path, what: str):
    if not path.is_file():
        raise ProtocolError(f"no {what} at {path}")
    try:
        return json.loads(path.read_text())
    except json.JSONDecodeError as e:
        raise ProtocolError(f"{path} is not valid JSON: {e}") from e


def load_methods(path: Path) -> dict[str, list[dict]]:
    """heading -> occurrences, in document order."""
    data = load_json(path, "methods file")
    if not isinstance(data, list) or not data:
        raise ProtocolError(f"{path} must be a non-empty list of sections")
    by_heading: dict[str, list[dict]] = {}
    for i, sec in enumerate(data):
        if not isinstance(sec, dict) or "heading" not in sec or "paras" not in sec:
            raise ProtocolError(f"{path}[{i}] needs both `heading` and `paras`")
        if not isinstance(sec["paras"], list):
            raise ProtocolError(f"{path}[{i}] `paras` must be a list of strings")
        by_heading.setdefault(sec["heading"], []).append(sec)
    return by_heading


def load_manifest(path: Path) -> dict:
    m = load_json(path, "manifest")
    for key in ("lab", "stamp", "protocols"):
        if key not in m:
            raise ProtocolError(f"{path} is missing required key `{key}`")
    if not re.fullmatch(r"[A-Za-z0-9]+", str(m["lab"])):
        raise ProtocolError(f"lab tag {m['lab']!r} must be alphanumeric")
    if not re.fullmatch(r"\d{6}", str(m["stamp"])):
        raise ProtocolError(f"stamp {m['stamp']!r} must be YYMMDD")
    if not isinstance(m["protocols"], list) or not m["protocols"]:
        raise ProtocolError(f"{path} declares no protocols")
    seen = set()
    for i, p in enumerate(m["protocols"]):
        for key in ("topic", "headings"):
            if key not in p:
                raise ProtocolError(f"{path} protocols[{i}] is missing `{key}`")
        if not TOPIC_RE.fullmatch(p["topic"]):
            raise ProtocolError(
                f"topic {p['topic']!r} must be CamelCase alphanumeric: it becomes "
                f"part of the filename and the SOP title")
        if p["topic"] in seen:
            raise ProtocolError(f"duplicate topic {p['topic']!r}")
        seen.add(p["topic"])
        if not p["headings"]:
            raise ProtocolError(f"protocols[{i}] ({p['topic']}) consumes no headings")
    return m


def filename_for(manifest: dict, topic: str) -> str:
    name = (f"P.{str(manifest['lab']).upper()}-{manifest['stamp']}"
            f"-V{int(manifest.get('version', 1))}_{topic}.docx")
    if not FILENAME_RE.fullmatch(name):
        raise ProtocolError(f"derived filename {name!r} is malformed")
    return name


# ── render ────────────────────────────────────────────────────────────────────

def new_document() -> "Document":
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        for side in ("top", "bottom", "left", "right"):
            setattr(section, f"{side}_margin", Inches(1))
    return doc


def render(doc: "Document", sections: list[dict]) -> int:
    """Italic heading then verbatim body, blank line between sections."""
    npara = 0
    for i, sec in enumerate(sections):
        if i:
            doc.add_paragraph()
        doc.add_paragraph().add_run(sec["heading"]).italic = True
        for text in sec["paras"]:
            doc.add_paragraph(text)
            npara += 1
    return npara


def verify_verbatim(path: Path, sections: list[dict]) -> list[str]:
    """Re-read the written file; return body paragraphs that drifted.

    Declared-non-verbatim sections are exempt: they are transcriptions by
    intent, and flagging them every run would train a curator to ignore the
    check that catches a real drift.
    """
    expected, exempt = [], set()
    for sec in sections:
        for text in sec["paras"]:
            expected.append(text)
            if sec.get("verbatim", True) is False:
                exempt.add(text)
    written = [p.text for p in Document(str(path)).paragraphs]
    drift = []
    for text in expected:
        if text in exempt:
            continue
        if text not in written:
            drift.append(text)
    return drift


# ── coverage ──────────────────────────────────────────────────────────────────

def edge_label(e: dict, nodes: dict | None = None) -> str:
    """`PARENT→CHILD`, suffixed `(count=0)` when the child tier has no rows.

    A protocol can document two edges at once, one built and one not. Flagging
    per edge rather than per row keeps that distinction: a flat ✅ on the row
    would claim the count=0 half is built, which is exactly the scope decision a
    reviewer needs to see.
    """
    label = f"`{e.get('from', '?')}→{e.get('to', '?')}`"
    if nodes is not None and nodes.get(e.get("to"), {}).get("count") == 0:
        label += " (count=0)"
    return label


def coverage(manifest: dict, tree: dict | None, sops: dict) -> str:
    lab = str(manifest["lab"]).upper()
    out = [f"# Protocol coverage: {lab}", ""]
    study = manifest.get("study")
    if study:
        out += [f"Protocols supporting {study}, checked against the assays that "
                f"license each parent→child edge in [`sample_tree.json`](../sample_tree.json).",
                ""]
    out += ["**Generated by `scripts/build_protocols.py`. Do not edit.** Change "
            "`_manifest.json` or `sample_tree.json` and re-run.", ""]

    if tree is None:
        out += ["No `sample_tree.json` was found, so neither table can be built. "
                "Run `/curate-sample-tree` first, then re-run this script.", ""]
        return "\n".join(out)

    nodes = {n["id"]: n for n in tree.get("nodes", []) if "id" in n}
    edges = [e for e in tree.get("edges", []) if "from" in e and "to" in e]

    def built(e: dict) -> bool | None:
        """True if the child tier has rows, False at count=0, None if unstated."""
        count = nodes.get(e["to"], {}).get("count")
        return None if count is None else count > 0

    edges_by_assay: dict[str, list[dict]] = {}
    for e in edges:
        for a in e.get("assays", []):
            edges_by_assay.setdefault(a, []).append(e)

    # Table A: protocol -> edge
    out += ["## Table A: Protocol → sample-tree edge", "",
            "| Protocol | SOP | Assay documented | Edge(s) in the sample tree | Note |",
            "|---|---|---|---|---|"]
    # Edges are dicts and several can compare equal, so index them positionally
    # rather than by value or by id().
    pos = {id(e): i for i, e in enumerate(edges)}
    documented_by: dict[int, list[str]] = {}
    referenced: set[str] = set()
    for p in manifest["protocols"]:
        fname = filename_for(manifest, p["topic"])
        assays = p.get("assays") or []
        hits: list[dict] = []
        for a in assays:
            for e in edges_by_assay.get(a, []):
                if all(pos[id(e)] != pos[id(h)] for h in hits):
                    hits.append(e)
        for e in hits:
            documented_by.setdefault(pos[id(e)], []).append(p["topic"])
            referenced.add(p["topic"])
        if not hits:
            sym, cells = "❌", "not referenced by any edge"
        else:
            sym = "⚠️" if all(built(e) is False for e in hits) else "✅"
            cells = ", ".join(edge_label(e, nodes) for e in hits)
        sop = sops.get(fname, {}).get("id", "n/a")
        out.append(f"| `{p['topic']}` | {sop} | {', '.join(assays) or 'n/a'} "
                   f"| {sym} {cells} | {p.get('note', '')} |")

    # Table B: edge -> protocol
    out += ["", "## Table B: Sample-tree edge → protocol coverage", "",
            f"All {len(edges)} edges in the current tree. A `(count=0)` edge is "
            "documented but produces no rows: a scope decision, not a "
            "documentation gap.", "",
            "| Edge | Assay | Protocol coverage |", "|---|---|---|"]
    for i, e in enumerate(edges):
        who = documented_by.get(i, [])
        cov = ("✅ " + ", ".join(f"`{t}`" for t in who)) if who else "❌ none"
        out.append(f"| {edge_label(e, nodes)} | {', '.join(e.get('assays', [])) or 'n/a'} "
                   f"| {cov} |")

    uncovered = [e for i, e in enumerate(edges) if not documented_by.get(i)]
    unreferenced = [p["topic"] for p in manifest["protocols"]
                    if p["topic"] not in referenced]
    out += ["", "## Gaps", ""]
    if uncovered:
        out.append(f"- **{len(uncovered)} edge(s) with no protocol:** "
                   + ", ".join(edge_label(e, nodes) for e in uncovered))
    else:
        out.append("- Every edge has protocol coverage.")
    if unreferenced:
        out.append(f"- **{len(unreferenced)} protocol(s) referenced by no edge:** "
                   + ", ".join(f"`{t}`" for t in unreferenced)
                   + ". Real methods with no row are expected when a tier is out of "
                     "scope; keep them registered so the record is complete.")
    out.append("")
    return "\n".join(out)


# ── main ──────────────────────────────────────────────────────────────────────

def build(args) -> int:
    root = Path(args.project_root).resolve()
    pdir = Path(args.protocols_dir) if args.protocols_dir else root / "protocols"
    manifest = load_manifest(Path(args.manifest) if args.manifest
                             else pdir / "_manifest.json")

    sops_path = pdir / "_sops.json"
    sops = json.loads(sops_path.read_text()) if sops_path.is_file() else {}

    tree_path = Path(args.sample_tree) if args.sample_tree else root / "sample_tree.json"
    tree = json.loads(tree_path.read_text()) if tree_path.is_file() else None
    if tree is None:
        print(f"  NOTE  no {tree_path}, so coverage tables will be empty")

    rc = 0
    if not args.coverage_only:
        methods = load_methods(Path(args.methods) if args.methods
                               else pdir / "_methods.json")
        used: dict[str, int] = {}
        wanted = {p["topic"] for p in manifest["protocols"]
                  if not args.only or args.only.lower() in p["topic"].lower()}
        if not wanted:
            raise ProtocolError(f"--only {args.only!r} matched no topic in the manifest")

        pdir.mkdir(parents=True, exist_ok=True)
        total = 0
        for p in manifest["protocols"]:
            # Walk EVERY protocol so the consumption counter stays honest under
            # --only; render only the ones asked for.
            sections = []
            for hd in p["headings"]:
                bucket = methods.get(hd)
                if not bucket:
                    raise ProtocolError(
                        f"protocol {p['topic']!r} wants heading {hd!r}, which is not "
                        f"in the methods file")
                n = used.get(hd, 0)
                sections.append(bucket[min(n, len(bucket) - 1)])
                used[hd] = n + 1
            if p["topic"] not in wanted:
                continue

            name = filename_for(manifest, p["topic"])
            target = pdir / name
            if target.exists() and not args.force:
                print(f"  SKIP  {name}  (already exists; pass --force to rewrite)")
                continue

            doc = new_document()
            npara = render(doc, sections)
            doc.save(str(target))
            drift = verify_verbatim(target, sections)
            total += npara
            flag = ""
            if drift:
                rc = 1
                flag = f"   {len(drift)} PARAGRAPH(S) NOT VERBATIM"
            print(f"  ✓ {name:<44} {len(p['headings']):>2} heading(s)  "
                  f"{npara:>2} para  {target.stat().st_size} bytes{flag}")
            for text in drift:
                print(f"        drifted: {text[:90]}…")

        for hd, bucket in methods.items():
            want, got = len(bucket), used.get(hd, 0)
            if want != got:
                rc = 1
                verb = "UNCONSUMED" if got < want else "OVERUSED"
                print(f"  {verb}: {hd!r} occurs {want}x in the methods file, "
                      f"consumed {got}x by the manifest")
        print(f"\n{len(wanted)} protocol(s) rendered, {total} body paragraphs")

    cov = pdir / "COVERAGE.md"
    pdir.mkdir(parents=True, exist_ok=True)
    cov.write_text(coverage(manifest, tree, sops))
    print(f"  ✓ {cov.relative_to(root) if cov.is_relative_to(root) else cov}")
    return rc


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    ap.add_argument("--project-root", default=".", help="curation project root")
    ap.add_argument("--protocols-dir", help="default <project-root>/protocols")
    ap.add_argument("--manifest", help="default <protocols-dir>/_manifest.json")
    ap.add_argument("--methods", help="default <protocols-dir>/_methods.json")
    ap.add_argument("--sample-tree", help="default <project-root>/sample_tree.json")
    ap.add_argument("--only", help="substring of a topic; render just that protocol")
    ap.add_argument("--force", action="store_true",
                    help="rewrite .docx files that already exist")
    ap.add_argument("--coverage-only", action="store_true",
                    help="regenerate COVERAGE.md without touching any .docx")
    args = ap.parse_args()
    try:
        sys.exit(build(args))
    except ProtocolError as e:
        sys.exit(f"protocols: {e}")


if __name__ == "__main__":
    main()
